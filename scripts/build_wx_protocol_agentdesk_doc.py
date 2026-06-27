from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/design/wx-protocol-agentdesk-qiyu-architecture.md"
OUT = Path("/Users/openclaw/Desktop/企微员工号-总部客服-门店知识库技术架构设计.docx")

INK = "111827"
MUTED = "4B5563"
BLUE = "2563EB"
TEAL = "0F766E"
GRAY_50 = "F8FAFC"
GRAY_100 = "F1F5F9"
GRAY_200 = "E2E8F0"


def set_run_font(run, size: float | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, size=9.2, color=INK)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after, bold in [
        ("Title", 21, INK, 0, 8, True),
        ("Subtitle", 10.5, MUTED, 0, 14, False),
        ("Heading 1", 15, BLUE, 14, 6, True),
        ("Heading 2", 12.2, TEAL, 9, 4, True),
        ("Heading 3", 11, INK, 7, 3, True),
        ("List Bullet", 9.8, INK, 0, 2, False),
        ("List Number", 9.8, INK, 0, 2, False),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.08

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"企微员工号 AgentDesk 架构设计 | {date.today().isoformat()}")
    set_run_font(run, size=8, color=MUTED)


def add_cover(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Title")
    run = p.add_run(title)
    set_run_font(run, size=21, color=INK)

    subtitle = doc.add_paragraph(style="Subtitle")
    run = subtitle.add_run("七鱼主链路已废弃；当前主链路为企微员工号、AgentDesk 网页工作台、门店知识库与总部网页客服。")
    set_run_font(run, size=10.5, color=MUTED)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("文档定位", "给新开发接手使用的技术架构、消息链路、页面设计与测试清单"),
        ("适用系统", "AgentDesk 二开版本，Go/Gin/GORM 后端 + Next.js/shadcn 后台"),
        ("核心入口", "企业微信员工号 / 微信协议 SAAS 回调"),
        ("核心原则", "所有消息先进 AgentDesk，再由 AgentDesk 决定 AI、人工、outbox 与审计"),
        ("版本", f"V2.0 / {date.today().isoformat()}"),
    ]
    for idx, (left, right) in enumerate(rows):
        shade_cell(table.rows[idx].cells[0], GRAY_100)
        shade_cell(table.rows[idx].cells[1], GRAY_50)
        set_cell_text(table.rows[idx].cells[0], left, bold=True)
        set_cell_text(table.rows[idx].cells[1], right)


def add_paragraph_with_inline_code(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part.strip("`"))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(TEAL)
        else:
            run = p.add_run(part)
            set_run_font(run)


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            if r_idx == 0:
                shade_cell(cell, GRAY_100)
            set_cell_text(cell, row[c_idx] if c_idx < len(row) else "", bold=r_idx == 0)


def build_doc() -> Path:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].lstrip("# ").strip()

    doc = Document()
    style_document(doc)
    add_cover(doc, title)

    table_buffer: list[list[str]] = []
    in_code = False
    code_lines: list[str] = []

    for raw in lines[1:]:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines).strip())
                run.font.name = "Consolas"
                run.font.size = Pt(8.2)
                run.font.color.rgb = RGBColor.from_string(MUTED)
                code_lines = []
                in_code = False
            else:
                flush_table(doc, table_buffer)
                table_buffer = []
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(set(cell) <= {"-", ":", " "} for cell in cells):
                continue
            table_buffer.append(cells)
            continue
        flush_table(doc, table_buffer)
        table_buffer = []

        if not line.strip():
            continue
        if line.startswith("## "):
            add_paragraph_with_inline_code(doc, line[3:].strip(), "Heading 1")
        elif line.startswith("### "):
            add_paragraph_with_inline_code(doc, line[4:].strip(), "Heading 2")
        elif line.startswith("#### "):
            add_paragraph_with_inline_code(doc, line[5:].strip(), "Heading 3")
        elif line.startswith("- "):
            add_paragraph_with_inline_code(doc, line[2:].strip(), "List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            # Keep explicit Markdown numbers so each checklist section starts at the source number.
            add_paragraph_with_inline_code(doc, line.strip())
        else:
            add_paragraph_with_inline_code(doc, line)

    flush_table(doc, table_buffer)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
