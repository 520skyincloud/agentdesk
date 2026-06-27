from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = Path("/Users/openclaw/Desktop/酒店AI店长客服系统技术架构与设计方案.docx")

INK = "111827"
MUTED = "4B5563"
BLUE = "2563EB"
DARK_BLUE = "1E3A8A"
TEAL = "0F766E"
GRAY_50 = "F8FAFC"
GRAY_100 = "F1F5F9"
GRAY_200 = "E2E8F0"
YELLOW_50 = "FFFBEB"
RED_50 = "FEF2F2"
GREEN_50 = "F0FDF4"


def set_run_font(run, font_name: str = "Microsoft YaHei", size: float | None = None) -> None:
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")


def set_paragraph_font(paragraph, font_name: str = "Microsoft YaHei") -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, bottom: int = 90, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: float = 9.5) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run, size=size)


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    style_specs = [
        ("Title", 22, INK, 0, 8, True),
        ("Subtitle", 10.5, MUTED, 0, 12, False),
        ("Heading 1", 15, BLUE, 15, 7, True),
        ("Heading 2", 12.5, DARK_BLUE, 10, 5, True),
        ("Heading 3", 11.5, TEAL, 7, 4, True),
        ("List Bullet", 10.2, INK, 0, 3, False),
        ("List Number", 10.2, INK, 0, 3, False),
    ]
    for name, size, color, before, after, bold in style_specs:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"酒店 AI 店长客服系统架构方案 | {date.today().isoformat()}")
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(run, size=8.5)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.add_run("酒店 AI 店长客服系统技术架构与设计方案")
    set_paragraph_font(title)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("基于 AgentDesk、企微接口 SAAS、门店知识矩阵与七鱼人工坐席的二开落地说明")
    set_paragraph_font(subtitle)

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [1.55, 4.95])
    rows = [
        ("文档定位", "给新开发快速接手使用的详细技术架构、设计架构与开发拆分说明"),
        ("当前系统", "AgentDesk：Go/Gin/GORM/simple 后端 + Next.js 后台 + MySQL + Qdrant"),
        ("目标链路", "企微员工号接口 SAAS / 小程序入口 -> AgentDesk -> AI/知识库 -> 七鱼人工 -> 原渠道回传"),
        ("一期目标", "先跑通 1 家门店与 1 个企微员工号，再扩到 100+ 门店多知识库、多实例、多路由"),
        ("版本", f"V1.0 / {date.today().isoformat()}"),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell_shading(meta.rows[i].cells[0], GRAY_100)
        set_cell_text(meta.rows[i].cells[0], k, bold=True)
        set_cell_text(meta.rows[i].cells[1], v)


def add_callout(doc: Document, title: str, body: str, fill: str = GRAY_50) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_run_font(r, size=10.5)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    set_paragraph_font(p2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        set_paragraph_font(p)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        set_paragraph_font(p)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], GRAY_100)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=INK, size=9.3)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, size=9.0)
    set_table_geometry(table, widths)
    doc.add_paragraph("")


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, GRAY_50)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.8)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_fonts.set(qn("w:ascii"), "Consolas")
    r_fonts.set(qn("w:hAnsi"), "Consolas")
    doc.add_paragraph("")


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. 一句话结论", level=1)
    add_callout(
        doc,
        "推荐架构",
        "不要让企微接口 SAAS、七鱼、小程序互相直接对接。AgentDesk 必须成为唯一会话中台：所有入口消息先进入 AgentDesk，"
        "统一写 Conversation/Message，再由路由状态决定交给 AI、七鱼人工或原渠道发送。这样才能解决多门店知识库、跨时间段上下文、人工回流和后续扩展问题。",
        GREEN_50,
    )
    add_bullets(
        doc,
        [
            "企微接口 SAAS：负责托管员工号实例、收外部联系人/群消息、以员工号身份发消息。",
            "AgentDesk：负责会话主数据、AI 回复、知识库选择、七鱼转人工、消息幂等、outbox 重试。",
            "七鱼：只作为人工坐席工作台，不直接面对企微客户；所有七鱼消息必须回流 AgentDesk。",
            "门店知识：公共知识库 + 门店知识矩阵 + 门店特殊知识库组合，按 storeId/brandLevel 动态选择。",
        ]
    )

    doc.add_heading("2. 当前系统现状", level=1)
    doc.add_paragraph("当前 AgentDesk 已具备客服会话中台雏形，二开时应复用已有基础，不要重写一套 IM 系统。")
    add_table(
        doc,
        ["已有能力", "当前代码/模型", "二开利用方式"],
        [
            ["会话主表", "models.Conversation", "继续作为所有入口的统一会话主表，新增路由状态字段或扩展表。"],
            ["消息主表", "models.Message", "所有客户、AI、人工、七鱼客服消息都写这里，保证上下文统一。"],
            ["渠道配置", "models.Channel", "新增/复用 channelType，配置企微接口 SAAS 实例、门店、AI Agent。"],
            ["异步投递", "models.ChannelMessageOutbox", "所有外部渠道下行消息先入 outbox，再由适配器发送，失败可重试。"],
            ["企微客服官方渠道", "WxWorkKFConversation / WxWorkKFMessageRef", "可参考其会话映射、幂等、发送失败处理，不建议和企微协议 SAAS 混用字段语义。"],
            ["企微 CLI/Hook 桥", "WxWorkCLIBridgeService / scripts/wecom-hook-bridge.mjs", "可作为企微协议 SAAS 适配器的雏形，后续改为云端回调模式。"],
            ["AI 运行时", "internal/ai/runtime", "客户消息写入后触发 AI，AI 回复写 Message，再进入 outbox。"],
            ["知识库", "Knowledge / AI Agent KnowledgeIDs / FastGPT 代理", "需要新增门店知识矩阵与动态知识选择层。"],
        ],
        [1.35, 1.95, 3.2],
    )

    doc.add_heading("3. 目标业务链路", level=1)
    add_code_block(
        doc,
        """用户入口
  ↓
NFC碰一碰 / 二维码 / 房卡 / 前台物料 / 微信小程序 / 企微员工号
  ↓
AgentDesk 接入层识别 channelId、storeId、customerExternalId
  ↓
ConversationService 创建或恢复会话
  ↓
知识路由：订单门店 > 入口门店 > 用户选择门店 > 未识别
  ↓
AI 店长回答
  ├─ 能回答：Message(AI) -> ChannelMessageOutbox -> 原入口发送
  └─ 不能回答 / 投诉 / 订单复杂：QiyuAdapter 转七鱼人工
          ↓
       七鱼客服回复回调 AgentDesk
          ↓
       Message(Agent/Qiyu) -> ChannelMessageOutbox -> 原入口发送
          ↓
       人工结束/超时/下班后，AI 带上下文继续接待"""
    )

    doc.add_heading("4. 总体技术架构", level=1)
    add_table(
        doc,
        ["层级", "模块", "职责", "技术落点"],
        [
            ["接入层", "小程序入口", "接收用户聊天、携带门店/订单/scene 信息。", "现有 API + 小程序新客服页。"],
            ["接入层", "企微接口 SAAS 适配器", "接收实例回调，发送员工号消息。", "新增 WxWorkProtocolAdapter。"],
            ["接入层", "七鱼适配器", "转人工、发客户消息到七鱼、接七鱼回调。", "新增 QiyuAdapterService。"],
            ["会话中台", "Conversation/Message", "统一存客户、AI、七鱼人工、系统事件消息。", "复用现有模型，补扩展表。"],
            ["路由层", "ConversationRouteService", "决定当前消息走 AI、七鱼还是人工等待。", "新增 service，避免散落 handler。"],
            ["AI 层", "AI 店长运行时", "知识检索、意图识别、回答、转人工判断。", "复用 runtime，补门店上下文。"],
            ["知识层", "门店知识矩阵", "同一问题模板，不同门店不同答案。", "新增 StoreKnowledgeMatrix。"],
            ["投递层", "ChannelMessageOutbox", "外部消息发送任务、重试、失败记录。", "复用并扩展 channelType。"],
            ["后台层", "配置与监控", "门店、实例、七鱼路由、知识矩阵、同步日志。", "Next.js dashboard 新页面。"],
        ],
        [0.8, 1.35, 2.25, 2.1],
    )

    doc.add_heading("5. 企微接口 SAAS 接入设计", level=1)
    doc.add_paragraph("根据目前确认的接口方式，企微员工号不部署在我们服务器上，而是由第三方 SAAS 托管实例。我们只接收回调并调用它的发送接口。")
    add_table(
        doc,
        ["设计项", "说明"],
        [
            ["实例模型", "1 个实例约等于 1 个登录员工号。100 家门店如果各用独立企微员工号，则约 100 个实例。"],
            ["通道模型", "建议 AgentDesk 中 1 个门店企微实例对应 1 个 Channel，channelType 可命名为 wxwork_protocol。"],
            ["回调入口", "POST /api/third/wxwork-protocol/callback，接收文本、图片、群消息、登录状态、掉线等事件。"],
            ["发送出口", "WxWorkProtocolOutboundWorker 轮询 outbox，调用 SAAS 发送文本/图片/卡片接口。"],
            ["会话标识", "以 instanceId + conversationId/externalUserId/groupId 作为外部会话唯一键。"],
            ["幂等键", "source=wxwork_protocol + instanceId + externalMsgId，落 MessageSyncLog 防止重复消费。"],
            ["掉线处理", "实例离线时写 ChannelHealthLog，后台告警；客户消息不丢，发送失败进入重试/人工告警。"],
        ],
        [1.45, 5.05],
    )

    doc.add_heading("6. 七鱼消息同步设计", level=1)
    add_callout(
        doc,
        "核心原则",
        "七鱼不能直接和企微互通。七鱼只和 AgentDesk 同步，AgentDesk 再把七鱼人工回复同步回客户原入口。这样才能保证消息历史、AI 上下文、跨时间段接管都统一。",
        YELLOW_50,
    )
    add_table(
        doc,
        ["方向", "动作", "AgentDesk 处理", "幂等要求"],
        [
            ["客户 -> 七鱼", "会话处于 QIYU_SERVING 时，客户新消息转发给七鱼。", "先写 Message(customer)，再调用 QiyuMessageSend。", "sourceMsgId 使用本地 messageId 或外部 msgId。"],
            ["七鱼 -> 客户", "七鱼客服回复回调 AgentDesk。", "先写 Message(agent/source=qiyu)，再入原渠道 outbox。", "qiyuMsgId 唯一，重复回调直接忽略。"],
            ["AI -> 客户", "AI 回复写 Message(ai)。", "按原入口 channelType 入 outbox。", "messageId + channelType 唯一。"],
            ["转人工摘要", "AI 触发转人工时发摘要给七鱼。", "包含门店、订单、客户最近 N 条消息、AI 判断原因。", "同一 handoffId 只发一次。"],
            ["人工结束", "七鱼结束/超时/客服离线回调。", "更新 QiyuConversation.status 和 ConversationRouteState。", "eventId 幂等。"],
        ],
        [0.9, 1.45, 2.45, 1.7],
    )

    doc.add_heading("7. 会话路由状态机", level=1)
    add_code_block(
        doc,
        """AI_SERVING
  - 默认状态。客户消息触发 AI。
  - AI 能答则直接回复。
  - 命中人工规则则进入 QIYU_PENDING。

QIYU_PENDING
  - 正在请求七鱼分配客服/客服组。
  - 成功：进入 QIYU_SERVING。
  - 失败：进入 AI_FALLBACK 或 WAIT_HUMAN。

QIYU_SERVING
  - 客户消息同步给七鱼。
  - 七鱼回复同步回原渠道。
  - AI 不主动回复，但可做摘要、质检、知识建议。

AI_FALLBACK
  - 七鱼不在线、排队过久、夜间服务。
  - AI 带七鱼历史上下文继续接待，并标记需人工跟进。

CLOSED
  - 会话关闭，不再自动回复；新消息按恢复策略重新打开。"""
    )

    doc.add_heading("8. 多门店知识库设计", level=1)
    doc.add_paragraph("酒店场景的特点不是每家店问题完全不同，而是问题模板固定、答案因门店不同而不同。因此应以“知识矩阵”而不是 100 个完全独立知识库为主。")
    add_table(
        doc,
        ["知识类型", "用途", "示例", "调用策略"],
        [
            ["公共知识库", "品牌、会员、发票、通用服务规则。", "AI店长是什么、会员权益、统一投诉流程。", "所有门店默认加载。"],
            ["门店知识矩阵", "固定问题模板，不同门店答案。", "早餐时间、停车场、洗衣房、健身房、前台电话。", "按 storeId 精确命中，优先级高。"],
            ["门店补充知识库", "门店特殊政策和临时通知。", "装修提醒、临时停水、附近路线。", "仅高阶品牌或已识别门店加载。"],
            ["订单上下文", "客户已入住/预订信息。", "入住日期、房型、订单所属门店。", "不进知识库，作为运行时上下文传给 AI。"],
        ],
        [1.15, 1.7, 1.8, 1.85],
    )
    add_code_block(
        doc,
        """知识选择伪代码：
if order.storeId exists:
    store = order.storeId
elif entry.scene.storeId exists:
    store = scene.storeId
elif user provided store:
    store = matchedStore
else:
    store = nil

knowledge = [company_common]
if store and brandLevel >= threshold:
    knowledge += [store_matrix(store), store_extra(store)]
elif store:
    knowledge += [store_matrix(store) only for highly deterministic FAQ]

answer = AI(question, orderContext, storeContext, knowledge)"""
    )

    doc.add_heading("9. 推荐新增数据模型", level=1)
    add_table(
        doc,
        ["模型", "核心字段", "说明"],
        [
            ["Store", "id、brandId、name、code、city、status", "门店主数据，100+ 门店的基础。"],
            ["Brand", "id、name、level、knowledgePolicy", "控制丽思及以上/以下品牌的知识调用策略。"],
            ["StoreEntryScene", "sceneCode、storeId、source、materialType", "NFC/二维码/房卡/前台物料入口映射。"],
            ["StoreKnowledgeItem", "storeId、questionKey、answer、effectiveAt、status", "门店知识矩阵；同一 questionKey 不同门店答案不同。"],
            ["WxWorkProtocolInstance", "channelId、instanceId、storeId、staffUserId、healthStatus", "企微接口 SAAS 实例与门店/通道绑定。"],
            ["ExternalConversationMap", "conversationId、channelType、externalConversationId、externalCustomerId", "统一外部会话映射，逐步替代混用 WxWorkKFConversation。"],
            ["QiyuConfig", "appKey、appSecret、callbackToken、status", "七鱼企业级配置。"],
            ["StoreQiyuRoute", "storeId、groupId、staffId、serviceTime、fallbackMode", "门店到七鱼客服组/客服路由。"],
            ["QiyuConversation", "conversationId、qiyuUid、sessionId、staffId、status、routeLockedUntil", "本地会话与七鱼会话映射。"],
            ["MessageSyncLog", "source、sourceMsgId、messageId、direction、syncStatus", "所有外部消息的幂等和排查依据。"],
            ["ConversationRouteState", "conversationId、routeStatus、routeTarget、handoffReason", "控制 AI/七鱼/等待人工状态机。"],
        ],
        [1.45, 2.15, 2.9],
    )

    doc.add_heading("10. 后端接口规划", level=1)
    add_table(
        doc,
        ["接口", "方法", "职责", "调用方"],
        [
            ["/api/third/wxwork-protocol/callback", "POST", "接收企微接口 SAAS 消息、实例状态、群消息。", "企微接口 SAAS"],
            ["/api/third/qiyu/callback", "POST", "接收七鱼客服消息、会话开始/结束/转接事件。", "七鱼"],
            ["/api/channel/config", "GET", "小程序/网页获取渠道配置。", "前端入口"],
            ["/api/im/conversation/init", "POST", "小程序初始化客户会话，携带 storeId/orderId/scene。", "小程序"],
            ["/api/dashboard/store/list", "ANY", "门店列表。", "后台"],
            ["/api/dashboard/store-knowledge/import", "POST", "批量导入门店 FAQ 矩阵。", "后台"],
            ["/api/dashboard/qiyu-route/update", "POST", "配置门店七鱼客服组/客服。", "后台"],
            ["/api/dashboard/channel/health", "GET", "查看企微实例、七鱼回调、outbox 健康状态。", "后台"],
        ],
        [2.0, 0.7, 2.55, 1.25],
    )

    doc.add_heading("11. 后端分层落地", level=1)
    add_table(
        doc,
        ["层", "新增内容", "注意事项"],
        [
            ["models", "新增 Store、QiyuConversation、MessageSyncLog 等实体。", "只定义字段和表映射，不写业务规则。"],
            ["repositories", "每个新模型生成基础 CRUD。", "只做数据访问，查询条件用 sqls.Cnd。"],
            ["services", "WxWorkProtocolService、QiyuAdapterService、ConversationRouteService、StoreKnowledgeService。", "业务编排、事务、状态机都在 service。"],
            ["handlers", "third 回调、dashboard 配置、api 会话入口。", "只解析参数、鉴权、调用 service、返回 httpx.WriteJSON。"],
            ["builders", "后台页面 response DTO 映射。", "禁止 DB 查询。"],
        ],
        [1.0, 2.55, 2.95],
    )

    doc.add_heading("12. 前端后台设计", level=1)
    add_table(
        doc,
        ["页面", "功能", "优先级"],
        [
            ["门店管理", "门店列表、品牌、城市、状态、入口码绑定。", "P0"],
            ["门店知识矩阵", "按问题模板维护每家店答案，支持 Excel 导入导出。", "P0"],
            ["渠道管理增强", "新增企微接口 SAAS 实例配置、健康状态、回调密钥。", "P0"],
            ["七鱼路由配置", "门店绑定七鱼客服组/客服、服务时间、兜底策略。", "P0"],
            ["会话监控", "显示 AI_SERVING/QIYU_SERVING/AI_FALLBACK 状态和同步失败。", "P1"],
            ["同步日志", "按 sourceMsgId、conversationId 查消息同步链路。", "P1"],
            ["运营报表", "AI解决率、转人工率、门店问题排行、七鱼响应时长。", "P2"],
        ],
        [1.45, 3.5, 0.8],
    )

    doc.add_heading("13. 部署架构", level=1)
    add_code_block(
        doc,
        """MVP 单机部署：
Linux Server 8C16G
  ├─ AgentDesk API / Dashboard
  ├─ MySQL
  ├─ Qdrant
  ├─ Redis 或内存队列（后续建议）
  └─ Worker: QiyuOutbound / WxWorkProtocolOutbound / Retry

外部服务：
  ├─ 企微接口 SAAS：员工号实例托管
  ├─ 七鱼：人工坐席
  ├─ FastGPT/云端知识库
  └─ DeepSeek/OpenAI/模型 API"""
    )
    add_table(
        doc,
        ["规模", "推荐配置", "说明"],
        [
            ["1-10 家试点", "4C8G / 100GB SSD / 5Mbps", "足够跑通链路，模型和企微实例走外部服务。"],
            ["100+ 门店一期", "8C16G / 200GB SSD / 10-20Mbps", "建议 MySQL 定期备份，outbox 和日志保留策略明确。"],
            ["规模化后", "应用、MySQL、Qdrant、Worker 拆机", "当消息量、知识库、报表压力上来再拆。"],
        ],
        [1.3, 2.15, 3.05],
    )

    doc.add_heading("14. 开发拆分计划", level=1)
    add_table(
        doc,
        ["阶段", "目标", "开发任务", "验收标准"],
        [
            ["第 0 阶段", "稳定现有环境", "确认 AgentDesk、模型配置、FastGPT、测试通道。", "本地/服务器可登录后台，AI 能回复。"],
            ["第 1 阶段", "企微接口 SAAS 接入", "新增 wxwork_protocol channel、回调、发送 worker、实例配置。", "私人微信给员工号发消息，AI 能自动回。"],
            ["第 2 阶段", "门店知识矩阵", "新增门店/品牌/FAQ 模板/答案导入，AI 动态选知识。", "同一问题在不同 storeId 下回答不同。"],
            ["第 3 阶段", "七鱼转人工", "新增 QiyuAdapter、QiyuConversation、七鱼回调和路由状态。", "AI 转人工后，七鱼回复能回到企微客户。"],
            ["第 4 阶段", "跨时间段接管", "routeLockedUntil、服务时间、超时回 AI、摘要上下文。", "夜间客户继续问，AI 带人工历史接待。"],
            ["第 5 阶段", "后台与监控", "同步日志、健康状态、失败重试、报表。", "运营能定位消息卡在哪一环。"],
        ],
        [0.9, 1.3, 2.75, 1.55],
    )

    doc.add_heading("15. 关键风险与处理", level=1)
    add_table(
        doc,
        ["风险", "影响", "处理策略"],
        [
            ["企微接口 SAAS 非官方", "账号风控、实例掉线、接口变更。", "合同确认 SLA；后台做健康监控；保留小程序/七鱼兜底。"],
            ["七鱼回调丢失或重复", "客户看不到人工回复或重复消息。", "MessageSyncLog 幂等；失败重试；按 qiyuMsgId 去重。"],
            ["多门店知识维护复杂", "答案混乱，AI 乱答。", "问题模板标准化，门店只填答案；上线前批量校验空值。"],
            ["AI 越权处理投诉/赔付", "服务风险。", "高风险意图强制转七鱼；AI 只做摘要不做承诺。"],
            ["两个外部通道抢 outbox", "消息发送错乱。", "一个 Channel 只能绑定一个 active sender；后台显示 sender owner。"],
            ["新开发不了解原项目分层", "代码难维护。", "严格按 models->repositories->services->handlers，禁止 handler 直查 DB。"],
        ],
        [1.55, 1.7, 3.25],
    )

    doc.add_heading("16. 新开发接手清单", level=1)
    add_numbered(
        doc,
        [
            "先跑通当前项目：docker compose up -d agent-desk，登录 /dashboard。",
            "阅读 models.Conversation、Message、Channel、ChannelMessageOutbox，理解会话中台基础。",
            "阅读 wxwork_cli_bridge_service.go，理解 inbound -> conversation -> message -> outbox 的现有样板。",
            "新增任何第三方接入时，先画清楚外部会话ID、本地 conversationId、messageId、sourceMsgId 四者关系。",
            "优先实现 service 层状态机和幂等，再做后台页面。",
            "每个阶段必须有端到端验收：客户发消息、AI/七鱼处理、原渠道收到回复、后台日志可查。",
        ],
    )

    doc.add_heading("17. 附录：建议命名", level=1)
    add_table(
        doc,
        ["类型", "建议名称"],
        [
            ["ChannelType", "wxwork_protocol、qiyu、mini_program"],
            ["Service", "WxWorkProtocolInboundService、WxWorkProtocolOutboundService、QiyuAdapterService、ConversationRouteService"],
            ["RouteStatus", "AI_SERVING、QIYU_PENDING、QIYU_SERVING、AI_FALLBACK、CLOSED"],
            ["Message source", "customer、ai、agent、qiyu、system"],
            ["Sync source", "wxwork_protocol、qiyu、mini_program"],
        ],
        [1.7, 4.8],
    )

    doc.add_paragraph("参考链接：")
    add_bullets(
        doc,
        [
            "企微接口 SAAS 文档：https://wework.apifox.cn/",
            "七鱼 OpenAPI 文档：https://qiyukf.apifox.cn/",
            "现有项目说明：/Users/openclaw/Downloads/agent-desk-main/AGENTS.md",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
