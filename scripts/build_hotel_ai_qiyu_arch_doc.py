from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "酒店AI店长与七鱼接入技术架构方案.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_east_asia(run, font_name: str = "Microsoft YaHei") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    set_run_east_asia(run)


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                break
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, val in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
                el = tc_mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    tc_mar.append(el)
                el.set(qn("w:w"), val)
                el.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 22, INK, 0, 8),
        ("Subtitle", 11, "666666", 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10
        if name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
            st.font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("酒店 AI 店长技术架构方案 | 2026-06-05")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("酒店 AI 店长与七鱼接入技术架构方案")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("面向多门店小程序接待、门店知识矩阵、七鱼人工坐席与私域承接的实施蓝图")

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1.65, 4.85])
    rows = [
        ("版本", "V0.1 / 方案讨论稿"),
        ("日期", "2026-06-05"),
        ("适用范围", "NFC/二维码/房卡/前台物料入口、小程序 AI 店长、七鱼人工客服、门店私域承接"),
        ("原则", "先做服务闭环和门店知识矩阵；企微私域先轻接入，重点门店再做深接入"),
    ]
    for idx, (k, v) in enumerate(rows):
        set_cell_shading(table.rows[idx].cells[0], LIGHT_GRAY)
        set_cell_text(table.rows[idx].cells[0], k, bold=True, color=INK)
        set_cell_text(table.rows[idx].cells[1], v)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(10.5)
    set_run_east_asia(r)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    set_table_geometry(table, widths)


def add_flow(doc: Document) -> None:
    doc.add_paragraph("主链路如下：")
    flow = (
        "用户触达入口\n"
        "  ↓\n"
        "NFC碰一碰 / 二维码 / 房卡 / 前台物料\n"
        "  ↓\n"
        "中间跳转页解析 sceneId / storeId / source\n"
        "  ↓\n"
        "微信小程序会话初始化\n"
        "  ↓\n"
        "AI店长接待\n"
        "  ↓\n"
        "识别订单与门店：订单门店 > 入口物料门店 > 用户选择门店 > 未识别\n"
        "  ↓\n"
        "判断品牌知识策略：高阶品牌 = 通用知识库 + 门店知识矩阵 + 门店补充知识库；低阶品牌 = 通用知识库\n"
        "  ↓\n"
        "知识匹配与回答判断\n"
        "  ├─ 能回答：AI店长直接回复\n"
        "  └─ 不能回答 / 投诉 / 高风险 / 订单复杂问题：转七鱼人工客服"
    )
    p = doc.add_paragraph()
    run = p.add_run(flow)
    run.font.name = "Consolas"
    set_run_east_asia(run)
    run.font.size = Pt(9.5)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. 目标与边界", level=1)
    add_callout(
        doc,
        "核心结论",
        "第一期应把 AgentDesk 二开为“会话中台 + AI 店长 + 门店知识矩阵 + 七鱼人工兜底 + 私域承接配置”。"
        "客户仍在小程序内聊天，七鱼只作为后端人工坐席；七鱼人工消息必须回流到 AgentDesk，保证跨时间段 AI 不失忆。",
    )
    add_bullets(
        doc,
        [
            "服务闭环：AI 能答则答，不能答时按门店/品牌路由到七鱼客服组。",
            "知识闭环：公共知识库负责品牌通用规则，门店知识矩阵负责同题不同答案，门店补充知识库处理特殊情况。",
            "私域闭环：第一期不强求 100 多家企微 API 深接入，而是按门店配置二维码/客服链接做轻私域承接。",
            "数据闭环：所有 AI、客户、七鱼人工消息统一沉淀在 AgentDesk，后续支持质检、报表和知识优化。",
        ],
    )

    doc.add_heading("2. 业务主链路", level=1)
    add_flow(doc)

    doc.add_heading("3. 总体系统架构", level=1)
    add_table(
        doc,
        ["模块", "职责", "第一期实现重点"],
        [
            ["入口层", "NFC、二维码、房卡、前台物料、中间跳转页、小程序入口。", "sceneId 解析、storeId/source 带入小程序会话。"],
            ["会话中台", "统一管理客户、会话、消息、路由状态和外部会话映射。", "AgentDesk Conversation/Message 扩展七鱼映射。"],
            ["AI 店长", "按门店和品牌知识策略回答客户问题，并判断是否转人工。", "接入门店知识矩阵和转人工触发规则。"],
            ["知识服务", "通用知识库、门店 FAQ 答案矩阵、门店补充知识库、品牌策略。", "同题不同答案用结构化矩阵，不把公共内容复制 100 份。"],
            ["七鱼适配器", "请求分配客服、转发客户消息、接收客服消息和会话事件。", "applyStaff、message/send、回调接收、sessionId 缓存。"],
            ["私域承接", "按门店提供企微二维码、客服链接、私域欢迎语和触发策略。", "轻接入优先，深接入只放重点门店或二期。"],
            ["后台管理", "门店、品牌、知识模板、七鱼路由、私域入口、服务时间。", "先满足运营配置和批量导入。"],
        ],
        [1.25, 2.85, 2.4],
    )

    doc.add_heading("4. 门店识别与知识调用策略", level=1)
    doc.add_paragraph("门店识别优先级必须稳定，避免客户无订单时 AI 反复追问门店。建议优先级如下：")
    add_numbered(
        doc,
        [
            "订单所属门店：客户有订单时，以订单门店为最高优先级。",
            "入口物料绑定门店：NFC、二维码、房卡、前台物料带 storeId 时，可直接识别门店。",
            "用户选择或输入门店：无订单且入口无法识别时，由 AI 主动询问并匹配门店。",
            "无法识别：仅调用通用知识库，并引导客户补充门店信息或转总客服。",
        ],
    )
    add_table(
        doc,
        ["场景", "知识范围", "说明"],
        [
            ["丽思及以上品牌且已识别门店", "通用知识库 + 门店知识矩阵 + 门店补充知识库", "门店答案优先，公共规则兜底。"],
            ["丽思以下品牌且已识别门店", "默认仅通用知识库", "可保留基础门店资料，但不启用完整门店知识。"],
            ["未识别门店", "仅通用知识库", "必要时询问门店或转总客服。"],
            ["投诉/高风险/订单复杂问题", "知识仅用于摘要，优先转人工", "不能让 AI 独立处理赔付、严重投诉等问题。"],
        ],
        [1.75, 2.75, 2.0],
    )

    doc.add_heading("5. 门店知识矩阵设计", level=1)
    doc.add_paragraph("你们的知识特点是“问题模板固定、答案按门店不同”。因此核心不应是 100 份散乱文档，而是模板化 FAQ 矩阵。")
    add_table(
        doc,
        ["层级", "数据形态", "例子"],
        [
            ["FAQ 模板", "标准问题、意图、标准回答模板、适用品牌。", "早餐几点？停车怎么停？能否提前入住？如何开发票？"],
            ["门店答案矩阵", "storeId + intent + answer + 生效状态。", "杭州西湖店 breakfast_time = 7:00-10:00。"],
            ["门店资料配置", "地址、电话、交通、早餐、停车、洗衣房、企微二维码。", "门店电话、停车入口、管家企微二维码。"],
            ["门店补充知识库", "临时公告、装修、施工、特殊活动、门店例外政策。", "本月停车场施工，请从东门进入。"],
            ["品牌公共知识库", "统一会员、发票、退改、服务标准、投诉处理框架。", "会员权益、品牌标准话术。"],
        ],
        [1.35, 2.65, 2.5],
    )

    doc.add_heading("6. 七鱼接入链路", level=1)
    doc.add_paragraph("七鱼接入采用“AgentDesk 中转”模式。客户不直接跳到七鱼界面，仍在小程序内聊天。")
    add_table(
        doc,
        ["步骤", "动作", "关键数据"],
        [
            ["1", "AI 判断不能回答或命中人工规则。", "conversationId、storeId、brandId、reason。"],
            ["2", "根据门店路由找到七鱼客服组/客服。", "qiyuGroupId、qiyuStaffId、serviceTime。"],
            ["3", "调用七鱼请求分配客服。", "applyStaff：uid、staffType=1、groupId、staffId。"],
            ["4", "缓存七鱼会话状态。", "sessionId、staffId、staffName、qiyuStatus。"],
            ["5", "把客户消息和转人工摘要发给七鱼。", "message/send：uid、msgType、content。"],
            ["6", "七鱼客服回复回调 AgentDesk。", "msgId、uid、staffId、content、timeStamp。"],
            ["7", "AgentDesk 写入消息并推给小程序。", "Message(senderType=agent, source=qiyu)。"],
            ["8", "会话结束/超时后 AI 可带上下文继续接待。", "handoffSummary、lastHumanReplyAt、routeLockedUntil。"],
        ],
        [0.55, 3.0, 2.95],
    )

    doc.add_page_break()
    doc.add_heading("7. 七鱼 API 与配置点", level=1)
    add_table(
        doc,
        ["用途", "接口/配置", "落地说明"],
        [
            ["请求分配客服", "POST https://qiyukf.com/openapi/event/applyStaff", "转人工时调用，优先按 staffId，其次 groupId。"],
            ["发送客户消息", "POST https://qiyukf.com/openapi/message/send", "客户后续消息直接转发，不重复分配客服。"],
            ["接收客服回复", "POST https://{QIYU_MSG_URL}", "七鱼后台“系统 → 扩展与集成 → 开发者ID”配置消息与事件接收 URL。"],
            ["会话开始通知", "POST https://{QIYU_MSG_URL}", "人工会话开始、转接时用于更新本地状态。"],
            ["用户资料", "/get_user_info", "可选；让七鱼侧边栏显示客户资料。"],
            ["订单信息", "/get_order", "可选；让七鱼侧边栏显示订单和门店信息。"],
        ],
        [1.3, 2.75, 2.45],
    )
    doc.add_paragraph("七鱼鉴权参数包括 appKey、time、checksum。checksum 规则为 SHA1(appSecret + md5 + time)，密钥必须只保存在服务端。")

    doc.add_page_break()
    doc.add_heading("8. 建议数据模型", level=1)
    add_table(
        doc,
        ["模型", "核心字段", "说明"],
        [
            ["Store", "id、name、brandId、city、status", "门店基础信息。"],
            ["Brand", "id、name、enableStoreKnowledge", "品牌层级和知识策略不要写死在代码里。"],
            ["StoreEntryScene", "sceneId、storeId、source、materialType", "NFC/二维码/房卡/物料与门店绑定。"],
            ["FAQTemplate", "intent、question、answerTemplate、brandScope", "固定问题模板。"],
            ["StoreFAQAnswer", "storeId、intent、answer、status、updatedAt", "每家门店同题不同答案。"],
            ["StorePrivateDomain", "storeId、qrUrl、contactLink、guideText", "轻私域承接配置。"],
            ["QiyuConfig", "appKey、appSecret、callbackToken、status", "七鱼企业级配置。"],
            ["StoreQiyuRoute", "storeId、groupId、staffId、serviceTime、fallbackTarget", "门店到七鱼客服组/客服的路由。"],
            ["QiyuConversation", "conversationId、uid、sessionId、staffId、status", "本地会话与七鱼会话映射。"],
            ["QiyuMessageRef", "messageId、qiyuMsgId、direction、rawPayload", "消息幂等、回放、排查。"],
        ],
        [1.55, 2.85, 2.1],
    )

    doc.add_heading("9. 转人工与跨时间段状态机", level=1)
    add_bullets(
        doc,
        [
            "新会话：按当前门店、品牌和服务时间决定是否允许转七鱼。",
            "已进入七鱼人工：不要因为服务时间到点立即切回 AI，应设置 routeLockedUntil 或等待七鱼会话结束。",
            "七鱼不在线 14005：AI 继续接待，必要时留工单或提示稍后由人工跟进。",
            "七鱼排队 14006：提示排队状态，也可让客户选择继续 AI 或退出排队。",
            "人工结束后：保留七鱼消息作为上下文，AI 后续接待要读取完整会话摘要。",
            "夜间新消息：如果七鱼不在服务时间，AI 接管，并带上白天七鱼人工会话上下文。",
        ],
    )

    doc.add_heading("10. 后台管理模块", level=1)
    add_table(
        doc,
        ["模块", "能力"],
        [
            ["门店管理", "门店基础信息、品牌归属、入口物料、服务时间、启停状态。"],
            ["知识模板管理", "统一维护固定问题模板和意图编码。"],
            ["门店答案矩阵", "按表格批量导入/导出 100 多家门店答案，并做缺失检查。"],
            ["七鱼配置", "appKey/appSecret、回调地址、客服组/客服 ID、连通性测试。"],
            ["路由规则", "按门店、品牌、时间、问题类型决定 AI/七鱼/工单/私域引导。"],
            ["私域配置", "门店企微二维码、客服链接、引导文案、触发条件。"],
            ["会话监控", "AI/七鱼状态、转人工原因、回调失败、消息重试。"],
        ],
        [1.55, 4.95],
    )

    doc.add_page_break()
    doc.add_heading("11. 实施阶段", level=1)
    add_table(
        doc,
        ["阶段", "范围", "交付结果"],
        [
            ["一期", "小程序入口、门店识别、AI 店长、门店知识矩阵、七鱼转人工与回流。", "形成服务闭环，100 多家门店可用。"],
            ["二期", "七鱼侧边栏用户/订单信息、夜间 AI 接管、质检基础、运营报表。", "人工和 AI 上下文统一，开始做效率管理。"],
            ["三期", "重点门店企微 API 深接入、自动私域标签、FastGPT/外部知识库扩展。", "做更深的私域运营和知识生态。"],
        ],
        [0.8, 3.4, 2.3],
    )

    doc.add_heading("12. 待确认事项", level=1)
    add_numbered(
        doc,
        [
            "小程序入口是否每个 NFC/二维码/房卡都能稳定带 sceneId 或 storeId。",
            "订单系统是否能通过用户身份或手机号查到订单、门店和品牌。",
            "品牌层级规则是否固定为“丽思及以上启用门店知识”，还是后台可配置。",
            "七鱼账号是否开通 OpenAPI、消息与事件回调、客服组 ID 查询能力。",
            "七鱼人工坐席是否全部留在七鱼工作台，不迁移到 AgentDesk。",
            "门店私域第一期是否采用二维码/客服链接轻承接，还是有重点门店需要企微 API 深接入。",
        ],
    )

    doc.add_heading("13. 参考接口资料", level=1)
    refs = [
        "七鱼调用流程：https://s.apifox.cn/apidoc/docs-site/605287/doc-437011",
        "请求分配客服：https://qiyukf.apifox.cn/api-10845470",
        "发送文本消息：https://qiyukf.apifox.cn/api-10845464",
        "接收七鱼消息：https://qiyukf.apifox.cn/api-10845467",
        "接收会话开始通知：https://qiyukf.apifox.cn/api-10845477",
        "第三方接口返回数据说明：https://qiyukf.apifox.cn/doc-437036",
    ]
    add_bullets(doc, refs)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
